import streamlit as st
import pdfplumber
import re
from langdetect import detect
from transformers import pipeline, MarianMTModel, MarianTokenizer
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from docx import Document

def limpiar_texto(texto):
    texto = re.sub(r'\[\d+\]', '', texto)
    texto = re.sub(r'\([A-Za-z, ]+\d{4}\)', '', texto)
    texto = re.sub(r'[^a-zA-Z0-9áéíóúñÁÉÍÓÚÑ\s\.,;:¡!¿?\-]', '', texto)
    texto = texto.lower()
    texto = re.sub(r'\n\s*\n', '\n', texto)
    texto = re.sub(r'\s+', ' ', texto).strip()
    return texto

st.title("📄 Herramienta de Resumen Automático de Artículos")

archivo_pdf = st.file_uploader("1️⃣ Sube tu archivo PDF aquí", type=["pdf"])

if archivo_pdf:
    with pdfplumber.open(archivo_pdf) as pdf:
        texto = ""
        for pagina in pdf.pages:
            texto += pagina.extract_text() + "\n"

    texto_limpio = limpiar_texto(texto)
    idioma = detect(texto_limpio)
    st.write(f"🌍 Idioma detectado: {idioma}")

    traducir = st.radio("2️⃣ ¿Deseas traducir el texto al español?", ["No", "Sí (si está en inglés)"])

    if traducir == "Sí (si está en inglés)" and idioma == "en":
        modelo_id = "Helsinki-NLP/opus-mt-en-es"
        tokenizer = MarianTokenizer.from_pretrained(modelo_id)
        model = MarianMTModel.from_pretrained(modelo_id)
        tokens = tokenizer(texto_limpio[:1000], return_tensors="pt", padding=True, truncation=True)
        traduccion = model.generate(**tokens)
        texto_limpio = tokenizer.decode(traduccion[0], skip_special_tokens=True)
        st.success("Texto traducido al español ✅")

    tipo = st.radio("3️⃣ ¿Qué tipo de resumen quieres?", ["Extractivo", "Abstractivo"])

    resumen = ""

    if st.button("📝 Generar resumen"):
        if tipo == "Extractivo":
            parser = PlaintextParser.from_string(texto_limpio, Tokenizer("spanish"))
            resumidor = TextRankSummarizer()
            resumen = "\n".join(str(oracion) for oracion in resumidor(parser.document, 5))
        else:
            resumen_ia = pipeline("summarization", model="facebook/bart-large-cnn")
            resumen = resumen_ia(texto_limpio[:1024], max_length=130, min_length=30, do_sample=False)[0]['summary_text']

        st.subheader("🧾 Resumen generado:")
        st.write(resumen)

        st.download_button("📥 Descargar como TXT", resumen, file_name="resumen.txt")

        doc = Document()
        doc.add_heading("Resumen generado", level=1)
        doc.add_paragraph(resumen)
        doc.save("resumen.docx")
        with open("resumen.docx", "rb") as f:
            st.download_button("📥 Descargar como DOCX", f, file_name="resumen.docx")
